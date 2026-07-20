import streamlit as st
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
from io import BytesIO

# --- DATABASE ---
REGIONI = ["dell'Abruzzo", "della Basilicata", "della Calabria", "della Campania", "dell'Emilia-Romagna", "del Friuli-Venezia Giulia", "del Lazio", "della Liguria", "della Lombardia", "delle Marche", "del Molise", "del Piemonte", "della Puglia", "della Sardegna", "della Sicilia", "della Toscana", "del Trentino-Alto Adige", "dell'Umbria", "della Valle d'Aosta", "del Veneto"]
CITTA = ["Agrigento", "Alessandria", "Ancona", "Aosta", "Arezzo", "Ascoli Piceno", "Asti", "Avellino", "Bari", "Barletta", "Belluno", "Benevento", "Bergamo", "Biella", "Bologna", "Bolzano", "Brescia", "Brindisi", "Cagliari", "Caltanissetta", "Campobasso", "Caserta", "Cassino", "Catania", "Catanzaro", "Chieti", "Como", "Cosenza", "Cremona", "Crotone", "Cuneo", "Enna", "Fermo", "Ferrara", "Firenze", "Foggia", "Forlì", "Frosinone", "Genova", "Gorizia", "Grosseto", "Imperia", "Isernia", "L'Aquila", "La Spezia", "Latina", "Lecce", "Lecco", "Livorno", "Lodi", "Lucca", "Macerata", "Mantova", "Massa Carrara", "Matera", "Messina", "Milano", "Modena", "Monza", "Napoli", "Novara", "Nuoro", "Oristano", "Padova", "Palermo", "Parma", "Pavia", "Perugia", "Pesaro", "Pescara", "Piacenza", "Pisa", "Pistoia", "Pordenone", "Potenza", "Prato", "Ragusa", "Ravenna", "Reggio Calabria", "Reggio Emilia", "Rieti", "Rimini", "Roma", "Rovigo", "Salerno", "Sassari", "Savona", "Siena", "Siracusa", "Sondrio", "Taranto", "Teramo", "Terni", "Torino", "Trani", "Trapani", "Trento", "Treviso", "Trieste", "Udine", "Varese", "Venezia", "Verbano-Cusio-Ossola", "Vercelli", "Verona", "Vibo Valentia", "Vicenza", "Viterbo"]

# --- FUNZIONE RESET MANUALE ---
def reset_procedura():
    # Testi
    st.session_state["rgr"] = ""
    st.session_state["sez"] = ""
    st.session_state["nome_cli"] = ""
    st.session_state["luogo_n"] = ""
    st.session_state["data_n"] = ""
    st.session_state["cf_cli"] = ""
    st.session_state["res_cli"] = ""
    # Numeri
    st.session_state["v_lite"] = 0.0
    st.session_state["cut_val"] = 0.0
    # Checkbox e altro
    st.session_state["u_check"] = False
    st.session_state["indet"] = False

def set_cell_shading(cell, color):
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)

def get_params(v):
    if v <= 1100: return (270.0, 270.0, 140.0, 340.0, "fino a € 1.100")
    elif v <= 5200: return (485.0, 405.0, 270.0, 610.0, "da € 1.101 a € 5.200")
    elif v <= 26000: return (1150.0, 875.0, 675.0, 1350.0, "da € 5.201 a € 26.000")
    elif v <= 52000: return (1960.0, 1285.0, 1080.0, 2365.0, "da € 26.001 a € 52.000")
    elif v <= 260000: return (3375.0, 2230.0, 1620.0, 3850.0, "da € 52.001 a € 260.000")
    else: return (5065.0, 3310.0, 2430.0, 5805.0, "da € 260.001 a € 520.000")

# --- INIZIALIZZAZIONE SESSION STATE ---
# Inizializziamo le chiavi se non esistono per evitare errori al primo avvio
if "rgr" not in st.session_state: reset_procedura()

# --- APP CONFIG ---
st.set_page_config(page_title="Nota Spese Pro", page_icon="⚖️")
st.title("⚖️ Generatore Nota Spese Tributaria")

# --- SIDEBAR ---
with st.sidebar:
    st.button("🗑️ PULISCI TUTTI I DATI", on_click=reset_procedura)
    st.divider()
    
    st.header("Procedimento")
    grado_sel = st.selectbox("Grado", ["I GRADO", "II GRADO"], key="grado_sel")
    
    if grado_sel == "I GRADO":
        sede = st.selectbox("Sede Corte (Città)", sorted(CITTA), key="sede_citta")
        g_h, g_c = f"DI PRIMO GRADO DI {sede.upper()}", f"di Primo Grado di {sede}"
    else:
        sede = st.selectbox("Sede Corte (Regione)", sorted(REGIONI), key="sede_regione")
        g_h, g_c = f"DI SECONDO GRADO {sede.upper()}", f"di Secondo Grado {sede}"

    st.text_input("Numero R.G.R.", key="rgr")
    st.text_input("Sezione", key="sez")
    st.checkbox("Includere data udienza?", key="u_check")
    
    u_data = ""
    if st.session_state["u_check"]:
        u_data = st.date_input("Data Udienza", datetime.date.today(), key="u_data").strftime("%d.%m.%Y")

    st.header("Valore e Spese")
    st.checkbox("Valore Indeterminato", key="indet")
    
    if st.session_state["indet"]:
        compl = st.selectbox("Complessità", ["Bassa", "Media", "Alta"], key="compl")
        v_calc = 25000.0 if compl == "Bassa" else 250000.0 if compl == "Media" else 500000.0
        val_txt = f"Valore Indeterminato (Complessità {compl.lower()})"
    else:
        st.number_input("Valore della lite (€)", min_value=0.0, key="v_lite")
        v_calc = float(st.session_state["v_lite"])
        val_txt = ""

    st.number_input("Contributo Unificato (C.U.T.) €", min_value=0.0, key="cut_val")

    st.header("Cliente")
    st.radio("Sesso", ["Femminile", "Maschile"], key="gender")
    st.text_input("Nome e Cognome", key="nome_cli")
    st.text_input("Luogo di Nascita", key="luogo_n")
    st.text_input("Data di Nascita", key="data_n")
    st.text_input("Codice Fiscale", key="cf_cli")
    st.text_input("Residenza", key="res_cli")

# --- LOGICA CALCOLO ---
prep = "della signora" if st.session_state["gender"] == "Femminile" else "del signor"
nasc = "nata" if st.session_state["gender"] == "Femminile" else "nato"

# --- GENERAZIONE DOCUMENTO ---
def create_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name, style.font.size = 'Calibri', Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # Intestazione
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_h = h.add_run(f"ALLA CORTE DI GIUSTIZIA TRIBUTARIA {g_h}")
    r_h.bold = True
    if st.session_state["sez"]: h.add_run(f"\nSEZIONE {st.session_state['sez'].upper()}").bold = True
    if u_data: h.add_run(f"\nUDIENZA DEL {u_data}").bold = True

    doc.add_paragraph("* * *").alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nota = doc.add_paragraph()
    p_nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nota.add_run("Nota spese ex art. 15, D.Lgs. 546/1992").bold = True
    doc.add_paragraph("* * *").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Corpo
    p_avv = doc.add_paragraph()
    p_avv.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_avv.add_run("Il ").bold = False
    p_avv.add_run("prof. dott. Mario Rovetti").bold = True
    p_avv.add_run(" ed il ").bold = False
    p_avv.add_run("dott. Sebastiano Barusco").bold = True
    p_avv.add_run(" con studio in Padova, via Cavazzana 5 (")
    contatti = p_avv.add_run("Barusco Rovetti & Associati, tel. 049-8752918, PEC: sebastiano.barusco@legalmail.it")
    contatti.underline = True
    p_avv.add_run(f"), in qualità di difensori {prep} ").bold = False
    p_avv.add_run(f"{st.session_state['nome_cli']}").bold = True
    p_avv.add_run(f", {nasc} a {st.session_state['luogo_n']} il {st.session_state['data_n']}, C.F. {st.session_state['cf_cli']}, residente in {st.session_state['res_cli']}")

    doc.add_paragraph("\nDEPOSITANO").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("la seguente nota spese:").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph(f"Competenza: Corte di Giustizia Tributaria {g_c}").space_after = Pt(12)

    # Calcoli
    p_par = get_params(v_calc)
    scl_final = val_txt if val_txt else p_par[4]
    tab_c = p_par[0] + p_par[1] + p_par[3]
    s_gen = tab_c * 0.15
    cpa = (tab_c + s_gen) * 0.04
    imp = tab_c + s_gen + cpa
    iva = imp * 0.22
    tot = imp + iva + st.session_state["cut_val"]

    # Tabella
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_r(l, v, bold=False, color=None):
        cells = table.add_row().cells
        cells[0].text, cells[1].text = l, v
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if bold:
            cells[0].paragraphs[0].runs[0].bold = True
            cells[1].paragraphs[0].runs[0].bold = True
        if color:
            set_cell_shading(cells[0], color)
            set_cell_shading(cells[1], color)

    add_r("VALORE DELLA CAUSA", scl_final, True, "EBEBEB")
    add_r("Fase di studio della controversia", f"€ {p_par[0]:,.2f}")
    add_r("Fase introduttiva del giudizio", f"€ {p_par[1]:,.2f}")
    add_r("Fase decisionale", f"€ {p_par[3]:,.2f}")
    add_r("COMPENSO TABELLARE", f"€ {tab_c:,.2f}", True, "F5F5F5")
    add_r("Rimborso spese generali (15%)", f"€ {s_gen:,.2f}")
    add_r("Contributo Cassa Previdenza (4%)", f"€ {cpa:,.2f}")
    add_r("TOTALE IMPONIBILE", f"€ {imp:,.2f}", True)
    add_r("I.V.A. (22%)", f"€ {iva:,.2f}")
    if st.session_state["cut_val"] > 0: add_r("Contributo Unificato Tributario (C.U.T.)", f"€ {st.session_state['cut_val']:,.2f}")
    add_r("IPOTESI DI COMPENSO LIQUIDABILE", f"€ {tot:,.2f}", True, "D9D9D9")

    # Firma
    mesi = {1:"Gennaio", 2:"Febbraio", 3:"Marzo", 4:"Aprile", 5:"Maggio", 6:"Giugno", 7:"Luglio", 8:"Agosto", 9:"Settembre", 10:"Ottobre", 11:"Novembre", 12:"Dicembre"}
    oggi = datetime.date.today()
    doc.add_paragraph(f"\nPadova, {oggi.day} {mesi[oggi.month]} {oggi.year}")
    sig = doc.add_paragraph("\nSebastiano Barusco")
    sig.add_run("\nFirmato digitalmente")

    out = BytesIO()
    doc.save(out)
    return out.getvalue()

# --- AZIONE GENERAZIONE ---
if st.button("🚀 GENERA DOCUMENTO"):
    try:
        st.download_button("📥 Scarica il documento Word", create_doc(), f"Nota_{st.session_state['nome_cli'].replace(' ','_')}.docx")
    except Exception as e:
        st.error(f"Errore: {e}")
